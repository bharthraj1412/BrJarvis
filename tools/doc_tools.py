# tools/doc_tools.py — BR JARVIS Word & PDF Document Generator Suite
"""
Automated Microsoft Word (.docx) and PDF (.pdf) Document Generator & Auto-Launcher.
"""
from __future__ import annotations

import json
import os
import re
import sys
import subprocess
from pathlib import Path
from typing import Any

from tools.registry import register_tool

try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    from fpdf import FPDF
    _FPDF_AVAILABLE = True
except ImportError:
    _FPDF_AVAILABLE = False


def _get_workspace_dir() -> Path:
    return Path(__file__).resolve().parent.parent


def _add_paragraph_runs(p, text: str):
    """Helper to split string on markdown '**' bold markers and append runs to a paragraph."""
    parts = text.split("**")
    is_bold = False
    for part in parts:
        if part:
            run = p.add_run(part)
            run.bold = is_bold
        is_bold = not is_bold


def _write_formatted_pdf_text(pdf, text: str):
    """Helper to split string on markdown '**' bold markers and write inline text to FPDF."""
    parts = text.split("**")
    is_bold = False
    for part in parts:
        if part:
            if is_bold:
                pdf.set_font("Helvetica", size=10, style="B")
            else:
                pdf.set_font("Helvetica", size=10)
            pdf.write(5, part)
        is_bold = not is_bold
    pdf.ln(6)


@register_tool(
    name="create_word_document",
    description="Create a formatted Microsoft Word (.docx) document with headers, styled paragraphs, tables, and auto-launch in Word.",
    parameters={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Document title"},
            "content": {"type": "string", "description": "Main document text content or markdown"},
            "filename": {"type": "string", "description": "Output filename ending in .docx"},
            "auto_open": {"type": "boolean", "description": "Whether to auto-launch Word"}
        },
        "required": ["title", "content"]
    }
)
def create_word_document(args: dict) -> str:
    """Create a Microsoft Word document (.docx)."""
    if not _DOCX_AVAILABLE:
        return "Error: 'python-docx' library is missing."

    title = args.get("title", "Document")
    content = args.get("content", "")
    filename = args.get("filename", "JARVIS_Document.docx")
    auto_open = args.get("auto_open", True)

    if not filename.endswith(".docx"):
        filename += ".docx"

    out_path = _get_workspace_dir() / filename

    doc = docx.Document()

    # Title
    h1 = doc.add_heading(title, level=0)
    h1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add content lines
    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        line_s = line.strip()
        if not line_s:
            i += 1
            continue

        # Parse markdown tables
        if line_s.startswith("|") and i + 1 < len(lines) and lines[i+1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            
            parsed_rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split("|")[1:-1]]
                is_sep = all(all(ch in "-: " for ch in cell) for cell in cells) if cells else False
                if not is_sep:
                    parsed_rows.append(cells)
            
            if parsed_rows:
                num_cols = max(len(row) for row in parsed_rows)
                table = doc.add_table(rows=0, cols=num_cols)
                try:
                    table.style = 'Light Shading Accent 1'
                except Exception:
                    table.style = 'Table Grid'
                
                for r_idx, row_cells in enumerate(parsed_rows):
                    row = table.add_row()
                    for c_idx, cell_text in enumerate(row_cells):
                        if c_idx < len(row.cells):
                            cell = row.cells[c_idx]
                            p = cell.paragraphs[0]
                            _add_paragraph_runs(p, cell_text)
                            if r_idx == 0:
                                for run in p.runs:
                                    run.bold = True
            continue

        # Headings
        if line_s.startswith("# "):
            doc.add_heading(line_s[2:], level=1)
        elif line_s.startswith("## "):
            doc.add_heading(line_s[3:], level=2)
        elif line_s.startswith("### "):
            doc.add_heading(line_s[4:], level=3)
        # Bullet list items
        elif line_s.startswith("- ") or line_s.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_paragraph_runs(p, line_s[2:])
        # Numbered list items
        elif re.match(r'^\d+\.\s+', line_s):
            match = re.match(r'^(\d+\.\s+)(.*)', line_s)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(match.group(1))
            run.bold = True
            _add_paragraph_runs(p, match.group(2))
        else:
            p = doc.add_paragraph()
            _add_paragraph_runs(p, line_s)
            
        i += 1

    try:
        doc.save(out_path)
    except PermissionError:
        import time
        ts = time.strftime("%H%M%S")
        out_path = _get_workspace_dir() / f"JARVIS_Document_{ts}.docx"
        doc.save(out_path)

    if auto_open and sys.platform == "win32":
        try:
            subprocess.Popen(f'start "" "{out_path}"', shell=True)
        except Exception:
            pass

    return f"⚡ Created Microsoft Word document: '{out_path}' and launched Word."


@register_tool(
    name="create_pdf_document",
    description="Create a formatted PDF (.pdf) document and auto-launch in default PDF viewer.",
    parameters={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Document title"},
            "content": {"type": "string", "description": "Main text content or markdown"},
            "filename": {"type": "string", "description": "Output filename ending in .pdf"},
            "auto_open": {"type": "boolean", "description": "Whether to auto-launch PDF viewer"}
        },
        "required": ["title", "content"]
    }
)
def create_pdf_document(args: dict) -> str:
    """Create a PDF document (.pdf)."""
    if not _FPDF_AVAILABLE:
        return "Error: 'fpdf' library is missing."

    title = args.get("title", "Document")
    content = args.get("content", "")
    filename = args.get("filename", "JARVIS_Document.pdf")
    auto_open = args.get("auto_open", True)

    if not filename.endswith(".pdf"):
        filename += ".pdf"

    out_path = _get_workspace_dir() / filename

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=16, style="B")
    pdf.cell(200, 10, txt=title.encode("latin-1", "replace").decode("latin-1"), ln=True, align="L")
    pdf.ln(5)

    pdf.set_font("Helvetica", size=10)

    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        line_clean = line.encode("latin-1", "replace").decode("latin-1").strip()
        if not line_clean:
            pdf.ln(3)
            i += 1
            continue

        # Parse markdown tables
        if line_clean.startswith("|") and i + 1 < len(lines) and lines[i+1].encode("latin-1", "replace").decode("latin-1").strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].encode("latin-1", "replace").decode("latin-1").strip().startswith("|"):
                table_lines.append(lines[i].encode("latin-1", "replace").decode("latin-1").strip())
                i += 1
            
            parsed_rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split("|")[1:-1]]
                is_sep = all(all(ch in "-: " for ch in cell) for cell in cells) if cells else False
                if not is_sep:
                    parsed_rows.append(cells)
            
            if parsed_rows:
                num_cols = len(parsed_rows[0])
                col_width = 190.0 / num_cols
                
                # Header row style
                pdf.set_font("Helvetica", size=9, style="B")
                pdf.set_fill_color(225, 235, 245)
                for cell in parsed_rows[0]:
                    pdf.cell(col_width, 8, cell, border=1, fill=True)
                pdf.ln()
                
                # Data rows style
                pdf.set_font("Helvetica", size=9)
                for row in parsed_rows[1:]:
                    for cell in row:
                        pdf.cell(col_width, 8, cell, border=1)
                    pdf.ln()
                pdf.ln(4)
            continue

        # Headings
        if line_clean.startswith("# "):
            pdf.ln(4)
            pdf.set_font("Helvetica", size=14, style="B")
            pdf.cell(0, 8, txt=line_clean[2:], ln=True)
            pdf.ln(2)
        elif line_clean.startswith("## "):
            pdf.ln(3)
            pdf.set_font("Helvetica", size=12, style="B")
            pdf.cell(0, 7, txt=line_clean[3:], ln=True)
            pdf.ln(2)
        elif line_clean.startswith("### "):
            pdf.ln(2)
            pdf.set_font("Helvetica", size=11, style="B")
            pdf.cell(0, 6, txt=line_clean[4:], ln=True)
            pdf.ln(1)
        # Bullet list items
        elif line_clean.startswith("- ") or line_clean.startswith("* "):
            pdf.set_font("Helvetica", size=10)
            pdf.write(5, "  o  ")
            _write_formatted_pdf_text(pdf, line_clean[2:])
        # Numbered list items
        elif re.match(r'^\d+\.\s+', line_clean):
            match = re.match(r'^(\d+\.\s+)(.*)', line_clean)
            pdf.set_font("Helvetica", size=10)
            pdf.write(5, f"  {match.group(1)}")
            _write_formatted_pdf_text(pdf, match.group(2))
        else:
            _write_formatted_pdf_text(pdf, line_clean)
            
        i += 1

    try:
        pdf.output(str(out_path))
    except PermissionError:
        import time
        ts = time.strftime("%H%M%S")
        out_path = _get_workspace_dir() / f"JARVIS_Document_{ts}.pdf"
        pdf.output(str(out_path))

    if auto_open and sys.platform == "win32":
        try:
            subprocess.Popen(f'start "" "{out_path}"', shell=True)
        except Exception:
            pass

    return f"⚡ Created PDF document: '{out_path}' and launched PDF reader."


@register_tool(
    name="generate_project_product_analysis",
    description="Generate a complete Product Analysis Report for B.R. JARVIS as Word (.docx) and PDF (.pdf) documents and auto-open them.",
    parameters={"type": "object", "properties": {}}
)
def generate_project_product_analysis(args: dict) -> str:
    """Generate complete Product Analysis report for B.R. JARVIS in Word & PDF formats."""
    doc_title = "Product Analysis Document: B.R. JARVIS"
    doc_text = """
# Product Analysis Document: B.R. JARVIS

## 1. Executive Summary
- **Product Name:** B.R. JARVIS (Advanced Agentic AI Operating System)
- **Identity:** Ultra-fast autonomous AI assistant designed for decisive multi-step reasoning, pair programming, and live desktop visual control.
- **Mission:** To eliminate conversational filler and deliver instant, high-precision software engineering actions.

## 2. Product Vision & Value Proposition
B.R. JARVIS shifts the paradigm from static autocomplete AI to an autonomous senior developer:
- **Zero-Filler Directive:** Delivers immediate, high-signal task resolution with minimal output tokens.
- **Local Gateway Integration:** Operates with unlimited request quotas via local gateway http://localhost:8045/v1.

## 3. Core Features & Capabilities
- **0-Token Intent Engine:** Executes common app launches, browser navigation, and diagnostics in 0ms with zero token cost.
- **Live OS Visual Controller:** Performs real-time desktop visual grounded automation (2160x1440 screen resolution).
- **Multi-Tab Excel & Document Generator:** Automatically creates formatted .xlsx, .docx, and .pdf analytical reports.

## 4. Architecture & Subsystems Inventory
- **Core Kernel:** Native C FNV-1a bridge, EventBus runtime, dependency injection container.
- **Action Suite:** Live OS Control, Computer Control, RAG Library, Desktop Automation.
- **Tool Registry:** 98+ registered tools across web, code, files, excel, process management, and audits.
- **UI HUD:** Tkinter glassmorphism display with Mission Control HUD and Max Control Center.

## 5. Security & Execution Safety
- Local-first workspace execution with absolute path validation.
- AST compilation checks and security vulnerability scanning.
"""
    res_word = create_word_document({"title": doc_title, "content": doc_text, "filename": "JARVIS_Product_Analysis.docx", "auto_open": True})
    res_pdf = create_pdf_document({"title": doc_title, "content": doc_text, "filename": "JARVIS_Product_Analysis.pdf", "auto_open": True})

    return f"{res_word}\n{res_pdf}"
