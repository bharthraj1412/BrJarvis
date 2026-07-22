import subprocess
import sys
import os

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

doc = docx.Document()
doc.add_heading('System Check', level=0)

table = doc.add_table(rows=2, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Status'

row_cells = table.rows[1].cells
row_cells[0].text = 'Core'
row_cells[1].text = 'Active'

save_path = os.path.expanduser('~/Master_Report.docx')
doc.save(save_path)
print(f'Successfully created {save_path}')